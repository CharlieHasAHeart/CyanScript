#!/usr/bin/env python3
import os
import re
import sys
import warnings
from typing import Optional

warnings.filterwarnings(
    "ignore",
    message="pkg_resources is deprecated as an API.*",
    category=UserWarning,
)

from docxtpl import DocxTemplate
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from markdown import markdown
from bs4 import BeautifulSoup, NavigableString

_HEADING_NUM_PATTERNS = [
    r"^\s*第\s*([0-9]+|[一二三四五六七八九十百千]+)\s*(章|节|部分|篇)\s*[:：、\.\s]*",
    r"^\s*[一二三四五六七八九十百千]+\s*[、\.\)]\s*",
    r"^\s*\d+(?:\.\d+)+\s*[\.\)]?\s*",
    r"^\s*\d+\s*[、\.\)]\s*",
]


def strip_heading_number(text: str) -> str:
    if not text:
        return text
    s = text.strip()
    for pat in _HEADING_NUM_PATTERNS:
        s_new = re.sub(pat, "", s).strip()
        if s_new != s:
            s = s_new
    return s


def update_fields_on_open(doc) -> None:
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        settings.append(update_fields)


def resolve_img_path(md_path: str, src: str) -> str:
    if os.path.isabs(src):
        return src
    base_dir = os.path.dirname(os.path.abspath(md_path))
    return os.path.join(base_dir, src)


def add_centered_image(subdoc, img_path: str, width_cm: float) -> None:
    p = subdoc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Cm(width_cm))


def add_caption(subdoc, text: str) -> None:
    if not text:
        return
    p = subdoc.add_paragraph(text)
    try:
        p.style = "图注"
    except KeyError:
        try:
            p.style = "Caption"
        except KeyError:
            p.style = "Normal"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_heading(subdoc, text: str, style_name: str) -> None:
    if not text:
        return
    p = subdoc.add_paragraph(text)
    try:
        p.style = style_name
    except KeyError:
        fallback_map = {
            "标题 1": "Heading 1",
            "标题 2": "Heading 2",
            "标题 3": "Heading 3",
            "标题 4": "Heading 4",
        }
        p.style = fallback_map.get(style_name, "Heading 1")
    fmt = p.paragraph_format
    fmt.left_indent = None
    fmt.first_line_indent = None
    fmt.hanging_indent = None


def add_paragraph(subdoc, text: str, style_name: str) -> None:
    if not text:
        return
    p = subdoc.add_paragraph(text)
    try:
        p.style = style_name
    except KeyError:
        fallback_map = {
            "正文": "Normal",
        }
        p.style = fallback_map.get(style_name, "Normal")


def add_paragraph_with_inline_code(subdoc, p_node, style_name: str) -> None:
    p = subdoc.add_paragraph()
    try:
        p.style = style_name
    except KeyError:
        p.style = "Normal"

    children = list(p_node.children)
    pending_thinspace = False

    for idx, child in enumerate(children):
        if isinstance(child, NavigableString):
            text = str(child)
            if pending_thinspace and text.startswith(" "):
                text = "\u2009" + text[1:]
                pending_thinspace = False
            elif pending_thinspace:
                pending_thinspace = False
            if text:
                if idx + 1 < len(children) and getattr(children[idx + 1], "name", None) == "code":
                    if text.endswith(" "):
                        text = text[:-1] + "\u2009"
                p.add_run(text)
        elif hasattr(child, "name") and child.name == "code":
            code_text = child.get_text()
            if code_text:
                run = p.add_run(code_text.replace(" ", "\u00A0"))
                try:
                    run.style = "行内代码"
                except KeyError:
                    pass
            if idx + 1 < len(children) and isinstance(children[idx + 1], NavigableString):
                pending_thinspace = True
        else:
            text = child.get_text(strip=False) if hasattr(child, "get_text") else str(child)
            if pending_thinspace and text.startswith(" "):
                text = "\u2009" + text[1:]
                pending_thinspace = False
            elif pending_thinspace:
                pending_thinspace = False
            if text:
                p.add_run(text)



def format_language(lang: str) -> str:
    if not lang:
        return ""
    mapping = {
        "py": "Python",
        "python": "Python",
        "js": "JavaScript",
        "javascript": "JavaScript",
        "ts": "TypeScript",
        "typescript": "TypeScript",
        "json": "JSON",
        "yaml": "YAML",
        "yml": "YAML",
        "bash": "Bash",
        "sh": "Shell",
        "shell": "Shell",
    }
    return mapping.get(lang.lower(), lang.capitalize())


def add_code_block(subdoc, pre_node) -> None:
    code_node = pre_node.find("code")
    lang = None

    if code_node and code_node.has_attr("class"):
        for cls in code_node["class"]:
            if cls.startswith("language-"):
                lang = cls.replace("language-", "").strip()
                break

    code_text = code_node.get_text() if code_node else pre_node.get_text()
    code_text = code_text.replace("\r\n", "\n").replace("\r", "\n").rstrip("\n")

    if lang:
        p_lang = subdoc.add_paragraph(f"语言：{format_language(lang)}")
        try:
            p_lang.style = "代码语言标记"
        except KeyError:
            p_lang.style = "代码块"

    for line in code_text.split("\n"):
        p = subdoc.add_paragraph(line)
        try:
            p.style = "代码块"
        except KeyError:
            p.style = "正文"


def add_table(subdoc, table_node) -> None:
    rows = []
    thead = table_node.find("thead")
    if thead:
        for tr in thead.find_all("tr", recursive=False):
            cells = [cell.get_text(strip=True) for cell in tr.find_all(["th", "td"], recursive=False)]
            if cells:
                rows.append(("header", cells))
    tbody = table_node.find("tbody")
    if tbody:
        for tr in tbody.find_all("tr", recursive=False):
            cells = [cell.get_text(strip=True) for cell in tr.find_all(["th", "td"], recursive=False)]
            if cells:
                rows.append(("body", cells))
    if not thead and not tbody:
        for tr in table_node.find_all("tr", recursive=False):
            cells = [cell.get_text(strip=True) for cell in tr.find_all(["th", "td"], recursive=False)]
            if cells:
                rows.append(("body", cells))
    if not rows:
        return

    max_cols = max(len(cells) for _, cells in rows)
    if max_cols == 0:
        return

    table = subdoc.add_table(rows=len(rows), cols=max_cols)
    try:
        table.style = "Table Grid"
    except KeyError:
        table.style = "Normal Table"

    for r_idx, (row_kind, cells) in enumerate(rows):
        row = table.rows[r_idx]
        for c_idx in range(max_cols):
            cell = row.cells[c_idx]
            text = cells[c_idx] if c_idx < len(cells) else ""
            cell.text = text
            style_name = "表格表头" if row_kind == "header" else "表格正文"
            for paragraph in cell.paragraphs:
                try:
                    paragraph.style = style_name
                except KeyError:
                    pass



def render_markdown_to_subdoc(subdoc, md_path: str) -> None:
    if not os.path.exists(md_path):
        print(f"[ERROR] Markdown not found: {md_path}")
        sys.exit(1)

    with open(md_path, "r", encoding="utf-8") as f:
        md_text = f.read()

    html = markdown(md_text, extensions=["extra"])
    soup = BeautifulSoup(html, "html.parser")
    body = soup.body if soup.body else soup

    fig_index = 0
    pending_table_caption = ""

    def handle_image(src: str, caption: str) -> None:
        nonlocal fig_index
        if not src:
            return
        name = caption.strip() if caption else ""
        if not name:
            base = os.path.basename(src)
            name = os.path.splitext(base)[0]
        fig_index += 1
        caption_text = f"图{fig_index} {name}" if name else f"图{fig_index}"
        img_path = resolve_img_path(md_path, src)
        if os.path.exists(img_path):
            add_centered_image(subdoc, img_path, 15)
            add_caption(subdoc, caption_text)
        else:
            add_paragraph(subdoc, f"[图片缺失: {src}]", "正文")
            add_caption(subdoc, caption_text)

    for node in body.children:
        if not hasattr(node, "name"):
            continue

        if node.name == "h1":
            add_heading(subdoc, strip_heading_number(node.get_text(strip=True)), "标题 1")
        elif node.name == "h2":
            add_heading(subdoc, strip_heading_number(node.get_text(strip=True)), "标题 2")
        elif node.name == "h3":
            add_heading(subdoc, strip_heading_number(node.get_text(strip=True)), "标题 3")
        elif node.name == "h4":
            add_heading(subdoc, strip_heading_number(node.get_text(strip=True)), "标题 4")
        elif node.name == "p":
            text = node.get_text(strip=True)
            if text:
                if re.match(r"^表\s*\d+\s+.+", text):
                    pending_table_caption = text
                else:
                    add_paragraph_with_inline_code(subdoc, node, "正文")
            for img in node.find_all("img"):
                handle_image(img.get("src", ""), img.get("alt", "") or "")
            for link in node.find_all("a"):
                href = link.get("href", "")
                if href.lower().endswith((".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp")):
                    handle_image(href, link.get_text(strip=True))
        elif node.name == "img":
            handle_image(node.get("src", ""), node.get("alt", "") or "")
        elif node.name == "table":
            if pending_table_caption:
                p = subdoc.add_paragraph(pending_table_caption)
                try:
                    p.style = "表注"
                except KeyError:
                    try:
                        p.style = "Caption"
                    except KeyError:
                        print("[WARN] 表注样式未找到，已使用默认正文样式。")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pending_table_caption = ""
            add_table(subdoc, node)
        elif node.name == "pre":
            add_code_block(subdoc, node)
        elif node.name == "a":
            href = node.get("href", "")
            if href.lower().endswith((".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp")):
                handle_image(href, node.get_text(strip=True))
            else:
                add_paragraph(subdoc, node.get_text(strip=True), "正文")


def prompt_input(label: str) -> str:
    return input(label).strip()


def safe_filename(value: str) -> str:
    cleaned = value.replace("\\", "_").replace("/", "_").strip()
    return cleaned or "output"


def resolve_template_path() -> str:
    env_path = os.getenv("CYANSCRIPT_TEMPLATE")
    candidates = []
    if env_path:
        if os.path.isabs(env_path):
            candidates.append(env_path)
        else:
            candidates.append(os.path.abspath(env_path))
    candidates.append(os.path.join(os.getcwd(), "assets", "reference.docx"))
    script_dir = os.path.dirname(os.path.abspath(__file__))
    candidates.append(os.path.join(script_dir, "assets", "reference.docx"))

    for path in candidates:
        if os.path.exists(path):
            return path
    print("[ERROR] template not found. Tried:")
    for path in candidates:
        print(f" - {path}")
    sys.exit(1)


def main() -> None:
    work_dir = prompt_input("工作目录(可选，留空则当前目录): ")
    if work_dir:
        work_dir = os.path.abspath(work_dir)
        if not os.path.isdir(work_dir):
            os.makedirs(work_dir, exist_ok=True)
        os.chdir(work_dir)

    template_path = resolve_template_path()

    software_name = prompt_input("软件名称: ")
    version = prompt_input("版本号: ")
    md_path = prompt_input("MD 文件: ")

    tpl = DocxTemplate(template_path)
    subdoc = tpl.new_subdoc()
    render_markdown_to_subdoc(subdoc, md_path)

    context = {
        "software_name": software_name,
        "version": version,
        "main_content": subdoc,
    }

    tpl.render(context)
    update_fields_on_open(tpl.docx)

    base_name = f"{safe_filename(software_name)}_{safe_filename(version)}_软件说明书.docx"
    final_path = os.path.join(os.getcwd(), base_name)

    tpl.save(final_path)
    print(f"[OK] Generated: {final_path}")


if __name__ == "__main__":
    main()
